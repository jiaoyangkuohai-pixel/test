from docx import Document
from docx.oxml.ns import qn
from io import BytesIO
from PIL import Image
from html import escape
from docx.document import Document as DocumentObject
from docx.table import Table, _Cell
from torch import return_types
 
def iter_block_items(parent: DocumentObject):
    """
    生成 `parent` 元素的每个块级元素，包括段落和表格。
    """
    parent_elm = parent.element.body
 
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            has_pic = bool(child.xpath(".//pic:pic") or child.xpath(".//w:drawing"))
            yield ('p_img' if has_pic else 'p'), child
        elif child.tag == qn('w:tbl'):
            yield 'tbl', child
        else:
            # print(f"unknown tag: {child.tag}")
            yield 'none', child
 



def get_picture(doc, img) -> list:
    """对图片做的一些操作"""
    embeds = img.xpath(".//a:blip/@r:embed")
    if not embeds:
        return []
    image_name = []
    for embed in embeds:
        related_part = doc.part.related_parts[embed]
        name = related_part.partname
        
        # image_name.append(f'<img src="{escape(str(name))}" />')
        image_name.append(f'{name}')
    #    image_blob = related_part.image.blob
    #    image = Image.open(BytesIO(image_blob))
        # 这里可以调用ocr模型对图片进行解释说明
    return image_name
 

def _get_grid_span(cell:_Cell):
    """返回单元格的水平合并跨度（colspan），默认 1。"""
    vals = cell._tc.xpath('.//w:tcPr/w:gridSpan/@w:val')
    if vals:
        try:
            return int(str(vals[0]))
        except Exception:
            return 1
    return 1



def _has_vmerge(cell:_Cell):
    """是否存在垂直合并标记（不关心 val）。"""
    return bool(cell._tc.xpath('.//w:tcPr/w:vMerge'))


def _cell_content_signature(doc, cell):
    """用于判等的内容签名：(纯文本, 图片路径列表)"""
    text_content = (cell.text or '').strip()
    img_paths = []
    imgs = cell._element.xpath('.//pic:pic')
    if imgs:
        for img in imgs:
            img_names = get_picture(doc, img)
            for name in img_names:
                img_paths.append(str(name))
    return (text_content, tuple(img_paths))

def _build_row_grid(row):
    """
    将一行按列网格展开，返回网格槽位列表。
    每个槽位包含：cell, tc_id, is_first_in_cell, colspan。
    """
    slots = []
    for cell in row.cells:
        colspan = _get_grid_span(cell)
        tc_id = id(cell._tc)
        # 第一个槽位标记为 true，其余为 false
        for i in range(colspan):
            slots.append({
                'cell': cell,
                'tc_id': tc_id,
                'is_first_in_cell': i == 0,
                'colspan': colspan,
            })
    return slots


def _parser_table(table: Table):
    html_lines = []
    html_lines.append('<table>')
    rows = table.rows

    # 预构建所有行的网格槽位
    row_grids = [_build_row_grid(row) for row in rows]

    for r_idx, row in enumerate(rows):
        html_lines.append('  <tr>')
        grid = row_grids[r_idx]

        consumed_first_positions = set()
        for g_idx, slot in enumerate(grid):
            if g_idx in consumed_first_positions:
                continue
            cell = slot['cell']
            if not slot['is_first_in_cell']:
                continue  # 属于某单元格的横向覆盖槽位，跳过

            # 若上一行同列也有相同内容签名且双方都有 vMerge，则当前为续接，跳过
            is_vmerge = _has_vmerge(cell)
            if r_idx > 0 and g_idx < len(row_grids[r_idx - 1]):
                prev_cell = row_grids[r_idx - 1][g_idx]['cell']
                if _has_vmerge(prev_cell) and is_vmerge:
                    if _cell_content_signature(doc, prev_cell) == _cell_content_signature(doc, cell):
                        continue

            # 计算横向合并：向右合并相邻、首槽且内容签名相同的单元格
            base_sig = _cell_content_signature(doc, cell)
            effective_colspan = slot['colspan']
            scan_pos = g_idx + effective_colspan
            merged_group_first_positions = []
            while scan_pos < len(grid):
                # 找到下一个首槽
                if not grid[scan_pos]['is_first_in_cell']:
                    scan_pos += 1
                    continue
                neighbor_slot = grid[scan_pos]
                neighbor_cell = neighbor_slot['cell']
                if _cell_content_signature(doc, neighbor_cell) == base_sig:
                    effective_colspan += neighbor_slot['colspan']
                    merged_group_first_positions.append(scan_pos)
                    scan_pos += neighbor_slot['colspan']
                else:
                    break

            # 这些被合并的首槽位后续不再单独渲染
            for pos in merged_group_first_positions:
                consumed_first_positions.add(pos)

            # 计算 rowspan：对合并组内每个首槽按同列规则计算，然后取最小值
            def compute_rowspan_at(col_index, base_cell):
                rs = 1
                if _has_vmerge(base_cell):
                    sig = _cell_content_signature(doc, base_cell)
                    rr = r_idx + 1
                    while rr < len(rows):
                        next_grid = row_grids[rr]
                        if col_index >= len(next_grid):
                            break
                        next_cell = next_grid[col_index]['cell']
                        if _has_vmerge(next_cell) and _cell_content_signature(doc, next_cell) == sig:
                            rs += 1
                            rr += 1
                        else:
                            break
                return rs

            rowspan_candidates = [compute_rowspan_at(g_idx, cell)]
            # 同时考虑合并组内其它首槽位置的纵向可合并高度
            accum = g_idx + slot['colspan']
            for pos in merged_group_first_positions:
                neighbor_slot = grid[pos]
                neighbor_cell = neighbor_slot['cell']
                rowspan_candidates.append(compute_rowspan_at(pos, neighbor_cell))
                accum = pos + neighbor_slot['colspan']
            rowspan = min(rowspan_candidates) if rowspan_candidates else 1

            colspan = effective_colspan

            # 内容：文本 + 图片
            cell_parts = []
            text_content = cell.text or ''
            if text_content:
                cell_parts.append(escape(text_content))
            imgs = cell._element.xpath(".//pic:pic")
            if imgs:
                for img in imgs:
                    img_names = get_picture(doc, img)
                    cell_parts.extend(img_names)
            cell_html = ' '.join(cell_parts) if cell_parts else ''

            attrs = []
            if rowspan > 1:
                attrs.append(f'rowspan="{rowspan}"')
            if colspan > 1:
                attrs.append(f'colspan="{colspan}"')
            attr_str = (' ' + ' '.join(attrs)) if attrs else ''
            html_lines.append(f'    <td{attr_str}>{cell_html}</td>')

        html_lines.append('  </tr>')
    html_lines.append('</table>')
    html_table = "\n".join(html_lines)
    # print(html_table)
    return html_table


def get_doc_text(doc: DocumentObject):
    paragraphs = doc.paragraphs
    tables = doc.tables
    paragraph_index = 0
    table_index = 0
    doc_parsed = []
    for block_type, block in iter_block_items(doc):
        if block_type == 'p':
            paragraph = paragraphs[paragraph_index]
            # print("Paragraph:", paragraph.text)
            doc_parsed.append(paragraph.text)
            paragraph_index += 1
        elif block_type == "p_img":
            # p_img 也是段落，需要同步推进段落索引以保持与文档块顺序一致
            paragraph = paragraphs[paragraph_index]
            # print("Paragraph:", paragraph.text)
            img_names = get_picture(doc, block)
            # print(f"Images: {img_names}")
            doc_parsed.append(paragraph.text)
            doc_parsed.extend(img_names)
            paragraph_index += 1
        elif block_type == 'tbl':
            table = tables[table_index]
            table_index += 1
            if table:
                htlm_table = _parser_table(table)
                doc_parsed.append(htlm_table)
    return doc_parsed
if __name__ == "__main__":
    # 打开文档
    doc_path = "/data/code/learn/work/chunk/data/ragflow_table_img.docx"
    doc = Document(doc_path)
    

    out = get_doc_text(doc)
    print("\n".join(out))
